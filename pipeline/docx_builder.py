"""Assemble the final markdown document into a .docx via a formatting template.

Markdown is the pipeline's internal artifact; ``.docx`` is one export target.
The mapping from document type to a title style lives in a small template
table, so a new output style is a data change, not new branching. python-docx
is imported lazily so the rest of the pipeline (and its tests) run without it.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class DocTemplate:
    """A minimal formatting template: heading text + base font size."""

    title_prefix: str
    body_font_pt: int = 12


# document_type -> template. Add a type by adding a row, not a code path.
TEMPLATES: dict[str, DocTemplate] = {
    "report": DocTemplate(title_prefix="Report", body_font_pt=12),
    "brief": DocTemplate(title_prefix="Brief", body_font_pt=11),
    "spec": DocTemplate(title_prefix="Specification", body_font_pt=11),
}


def _iter_blocks(markdown: str):
    """Yield ('h1'|'h2'|'p', text) blocks from a small markdown subset."""
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            yield "h2", line[3:].strip()
        elif line.startswith("# "):
            yield "h1", line[2:].strip()
        else:
            yield "p", line.strip()


def build_docx(markdown: str, out_path: Path, document_type: str = "report", title: str = "") -> Path:
    """Render ``markdown`` to a .docx at ``out_path`` using the type template."""
    from docx import Document  # lazy import
    from docx.shared import Pt

    template = TEMPLATES.get(document_type, TEMPLATES["report"])
    doc = Document()
    doc.add_heading(title or f"{template.title_prefix}", level=0)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(template.body_font_pt)

    for kind, text in _iter_blocks(markdown):
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
